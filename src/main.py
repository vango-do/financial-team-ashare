from __future__ import annotations

import json
import os
import re
import sys
from datetime import datetime
from html import unescape
from pathlib import Path

from colorama import Fore, Style, init
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from dotenv import load_dotenv
from langchain_core.messages import HumanMessage
from langgraph.graph import END, StateGraph

os.environ.setdefault("TQDM_DISABLE", "1")

from src.agents.portfolio_manager import portfolio_management_agent
from src.agents.risk_manager import risk_management_agent
from src.cli.input import parse_cli_inputs
from src.graph.state import AgentState
from src.rag.master_retrieval import isolation_audit
from src.utils.analysts import ANALYST_ORDER, get_agent_display_name, get_analyst_nodes
from src.utils.display import print_trading_output
from src.utils.progress import progress

try:
    import akshare as ak
except Exception:  # pragma: no cover
    ak = None

load_dotenv()
init(autoreset=True)

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

PROJECT_ROOT = Path(__file__).resolve().parents[1]


def parse_hedge_fund_response(response):
    try:
        return json.loads(response)
    except Exception:
        return None


def start(state: AgentState):
    return state


def create_workflow(selected_analysts=None):
    workflow = StateGraph(AgentState)
    workflow.add_node("start_node", start)

    analyst_nodes = get_analyst_nodes()
    if selected_analysts is None:
        selected_analysts = list(analyst_nodes.keys())

    for analyst_key in selected_analysts:
        node_name, node_func = analyst_nodes[analyst_key]
        workflow.add_node(node_name, node_func)
        workflow.add_edge("start_node", node_name)

    workflow.add_node("risk_management_agent", risk_management_agent)
    workflow.add_node("portfolio_manager", portfolio_management_agent)

    for analyst_key in selected_analysts:
        node_name = analyst_nodes[analyst_key][0]
        workflow.add_edge(node_name, "risk_management_agent")

    workflow.add_edge("risk_management_agent", "portfolio_manager")
    workflow.add_edge("portfolio_manager", END)
    workflow.set_entry_point("start_node")
    return workflow


def run_hedge_fund(
    tickers: list[str],
    start_date: str,
    end_date: str,
    portfolio: dict,
    show_reasoning: bool = False,
    selected_analysts: list[str] | None = None,
    model_name: str = "deepseek-ai/DeepSeek-V3",
    model_provider: str = "DeepSeek",
):
    progress.start()
    try:
        workflow = create_workflow(selected_analysts if selected_analysts else None)
        agent = workflow.compile()
        final_state = agent.invoke(
            {
                "messages": [
                    HumanMessage(
                        content="\u8bf7\u57fa\u4e8e\u5df2\u63d0\u4f9b\u6570\u636e\u5b8c\u6210 A \u80a1\u4ea4\u6613\u51b3\u7b56\u3002"
                    )
                ],
                "data": {
                    "tickers": tickers,
                    "portfolio": portfolio,
                    "start_date": start_date,
                    "end_date": end_date,
                    "analyst_signals": {},
                },
                "metadata": {
                    "show_reasoning": show_reasoning,
                    "model_name": model_name,
                    "model_provider": model_provider,
                },
            }
        )
        return {
            "decisions": parse_hedge_fund_response(final_state["messages"][-1].content),
            "analyst_signals": final_state["data"]["analyst_signals"],
            "retrieval_logs": final_state["data"].get("retrieval_logs", []),
            "chief_memory_records": final_state["data"].get("chief_memory_records", []),
        }
    finally:
        progress.stop()


def _agent_sort_key(agent_node: str) -> int:
    analyst_order = {analyst_key: idx for idx, (_, analyst_key) in enumerate(ANALYST_ORDER)}
    if agent_node == "risk_management_agent":
        return len(analyst_order) + 1
    if agent_node.endswith("_agent"):
        return analyst_order.get(agent_node[: -len("_agent")], 999)
    return analyst_order.get(agent_node, 999)


def _safe_file_name(name: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "_", str(name or "").strip())
    return cleaned or "A\u80a1\u6807\u7684"


def _resolve_runtime_path(env_key: str, default_relative: str) -> Path:
    raw_value = os.getenv(env_key, default_relative)
    path = Path(raw_value).expanduser()
    if not path.is_absolute():
        path = (PROJECT_ROOT / path).resolve()
    return path


def _safe_display_path(path: Path | None) -> str:
    if path is None:
        return ""
    try:
        rel = path.resolve().relative_to(PROJECT_ROOT.resolve())
        return f"./{rel.as_posix()}"
    except Exception:
        return path.name


def _get_desktop_dir() -> Path | None:
    custom = os.getenv("DESKTOP_OUTPUT_DIR")
    if not custom:
        return None
    path = Path(custom).expanduser()
    if not path.is_absolute():
        path = (PROJECT_ROOT / path).resolve()
    return path


def _resolve_stock_name(ticker: str) -> str:
    if ak is None:
        return ticker
    try:
        df = ak.stock_individual_info_em(symbol=ticker)
        if df is None or df.empty:
            return ticker
        for _, row in df.iterrows():
            k = str(row.iloc[0]).strip()
            v = str(row.iloc[1]).strip()
            if any(x in k for x in ("\u80a1\u7968\u7b80\u79f0", "\u7b80\u79f0", "\u540d\u79f0")) and v:
                return v
    except Exception:
        return ticker
    return ticker


def _signal_cn(signal: str) -> str:
    mapping = {"BULLISH": "\u770b\u591a", "BEARISH": "\u770b\u7a7a", "NEUTRAL": "\u4e2d\u6027"}
    return mapping.get((signal or "").upper(), signal or "")


def _action_cn(action: str) -> str:
    mapping = {
        "BUY": "\u4e70\u5165",
        "SELL": "\u5356\u51fa",
        "HOLD": "\u6301\u6709",
        "SHORT": "\u505a\u7a7a",
        "COVER": "\u56de\u8865",
    }
    return mapping.get((action or "").upper(), action or "")


def _normalize_reasoning(reasoning) -> str:
    if reasoning is None:
        return ""
    if isinstance(reasoning, str):
        text = reasoning.strip()
    elif isinstance(reasoning, (dict, list)):
        text = json.dumps(reasoning, ensure_ascii=False, indent=2)
    else:
        text = str(reasoning)

    alpha_count = len(re.findall(r"[A-Za-z]", text))
    han_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    if han_count == 0 and alpha_count > 0:
        return "\u4e2d\u6587\u8f93\u51fa\u7ea6\u675f\u89e6\u53d1\uff1a\u68c0\u6d4b\u5230\u82f1\u6587\u5185\u5bb9\uff0c\u5df2\u964d\u7ea7\u5c55\u793a\u3002"
    return text


def _safe_confidence(value: object) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    try:
        return float(str(value).replace("%", "").strip())
    except Exception:
        return 0.0


def _latest_retrieval_log(retrieval_logs: list[dict], agent: str, ticker: str) -> dict:
    for row in reversed(retrieval_logs):
        if row.get("agent") != agent:
            continue
        query_text = str(row.get("query") or "")
        if ticker in query_text or not ticker:
            return row
    return {}


def _extract_reasoning_parts(reasoning: object, retrieval_log: dict) -> dict[str, str | list[str]]:
    evidence_ids: list[str] = []
    if isinstance(reasoning, dict):
        evidence_ids = list(reasoning.get("evidence_ids") or [])
    if not evidence_ids:
        evidence_ids = list(retrieval_log.get("hit_ids") or [])

    if isinstance(reasoning, dict):
        conclusion = str(reasoning.get("conclusion") or reasoning.get("signal_determination") or "").strip()
        basis = str(reasoning.get("basis") or reasoning.get("details") or "").strip()
        risk = str(reasoning.get("risk") or "").strip()
        trigger = str(reasoning.get("trigger") or "").strip()

        if not basis:
            basis = json.dumps(reasoning, ensure_ascii=False, indent=2)
        if not conclusion:
            conclusion = "\u89c1\u4f9d\u636e"
        if not risk:
            risk = "\u82e5\u5173\u952e\u8d22\u52a1\u3001\u60c5\u7eea\u6216\u653f\u7b56\u8bc1\u636e\u53d8\u5316\uff0c\u539f\u7ed3\u8bba\u53ef\u80fd\u5931\u6548\u3002"
        if not trigger:
            trigger = "\u5f53\u4f30\u503c\u3001\u8d44\u91d1\u9762\u6216\u653f\u7b56\u51fa\u73b0\u53cd\u5411\u53d8\u5316\u65f6\u89e6\u53d1\u590d\u8bc4\u3002"

        return {
            "conclusion": conclusion,
            "basis": basis,
            "risk": risk,
            "trigger": trigger,
            "evidence_ids": evidence_ids,
        }

    text = _normalize_reasoning(reasoning)
    return {
        "conclusion": text[:180] if text else "\u8bc1\u636e\u4e0d\u8db3",
        "basis": text or "\u8bc1\u636e\u4e0d\u8db3",
        "risk": "\u82e5\u6838\u5fc3\u5047\u8bbe\u88ab\u8bc1\u4f2a\uff0c\u9700\u8981\u53ca\u65f6\u64a4\u9500\u8be5\u89c2\u70b9\u3002",
        "trigger": "\u4ef7\u683c\u3001\u8d22\u62a5\u3001\u653f\u7b56\u6216\u8d44\u91d1\u9762\u51fa\u73b0\u7a81\u53d8\u65f6\u7acb\u5373\u590d\u8bc4\u3002",
        "evidence_ids": evidence_ids,
    }


def _majority_signal(signals: list[str]) -> str:
    if not signals:
        return "NEUTRAL"
    counts: dict[str, int] = {}
    for signal in signals:
        key = str(signal or "").upper()
        counts[key] = counts.get(key, 0) + 1
    return max(counts, key=counts.get)


def _append_quote_block(lines: list[str], text: str) -> None:
    normalized = _normalize_reasoning(text)
    if not normalized:
        lines.append("> \uff08\u65e0\uff09")
        return
    for raw_line in normalized.splitlines():
        lines.append(f"> {raw_line}" if raw_line else ">")


def _build_consolidated_markdown(result: dict, tickers: list[str], start_date: str, end_date: str) -> str:
    decisions = result.get("decisions", {}) or {}
    analyst_signals = result.get("analyst_signals", {}) or {}
    retrieval_logs = result.get("retrieval_logs", []) or []
    audit = isolation_audit(retrieval_logs)
    now_text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    ticker_names = [f"{_resolve_stock_name(t)}\uff08{t}\uff09" for t in tickers]

    def _h2(text: str) -> str:
        return f'<div style="font-size:15pt;font-weight:700;">{text}</div>'

    def _h4(text: str) -> str:
        return f'<div style="font-size:14pt;font-weight:700;">{text}</div>'

    lines: list[str] = [
        '<div style="text-align:center;font-size:26pt;font-weight:700;">A股多Agent融合综合决策报告</div>',
        "",
        f"\u751f\u6210\u65f6\u95f4\uff1a{now_text}",
        f"\u5206\u6790\u533a\u95f4\uff1a{start_date} \u81f3 {end_date}",
        f"\u6807\u7684\u5217\u8868\uff1a{'\u3001'.join(ticker_names)}",
        f"\u53c2\u4e0e\u5927\u5e08\uff1a{'\u3001'.join(get_agent_display_name(f'{k}_agent') for _, k in ANALYST_ORDER)}",
        "",
    ]

    portfolio_rows: list[dict] = []
    conflict_records: list[dict] = []
    risk_items: list[str] = []
    invalidation_items: list[str] = []

    for ticker in tickers:
        stock_name = _resolve_stock_name(ticker)
        decision = decisions.get(ticker, {}) or {}

        analyst_entries: list[dict] = []
        signal_vector: list[str] = []
        for agent, by_ticker in sorted(analyst_signals.items(), key=lambda x: _agent_sort_key(x[0])):
            if agent == "risk_management_agent" or ticker not in by_ticker:
                continue
            signal_obj = by_ticker.get(ticker, {}) or {}
            signal = str(signal_obj.get("signal", "")).upper()
            confidence = _safe_confidence(signal_obj.get("confidence", 0))
            confidence_text = f"{confidence:.2f}%"
            retrieval_log = _latest_retrieval_log(retrieval_logs, agent, ticker)
            reasoning_parts = _extract_reasoning_parts(signal_obj.get("reasoning"), retrieval_log)
            evidence_ids = [str(i) for i in reasoning_parts.get("evidence_ids", []) if str(i).strip()]
            evidence_text = "\u3001".join(evidence_ids) if evidence_ids else "\u65e0"

            signal_vector.append(signal)
            analyst_entries.append(
                {
                    "role": get_agent_display_name(agent),
                    "signal": signal,
                    "signal_cn": _signal_cn(signal),
                    "confidence_text": confidence_text,
                    "conclusion": str(reasoning_parts.get("conclusion") or "").strip(),
                    "basis": str(reasoning_parts.get("basis") or "").strip(),
                    "risk": str(reasoning_parts.get("risk") or "").strip(),
                    "trigger": str(reasoning_parts.get("trigger") or "").strip(),
                    "evidence_text": evidence_text,
                }
            )
            risk_items.append(str(reasoning_parts.get("risk") or ""))
            invalidation_items.append(str(reasoning_parts.get("trigger") or ""))

        majority = _majority_signal(signal_vector)
        majority_cn = _signal_cn(majority)
        action = str(decision.get("action", "")).upper()
        confidence = decision.get("confidence", 0)
        confidence_text = f"{confidence:.1f}%" if isinstance(confidence, (int, float)) else f"{confidence}%"
        decision_reasoning = _normalize_reasoning(decision.get("reasoning"))

        bullish = sum(1 for s in signal_vector if s == "BULLISH")
        bearish = sum(1 for s in signal_vector if s == "BEARISH")
        neutral = sum(1 for s in signal_vector if s == "NEUTRAL")
        portfolio_rows.append(
            {
                "ticker": ticker,
                "action": _action_cn(action),
                "quantity": decision.get("quantity", 0),
                "confidence_text": confidence_text,
                "bullish": bullish,
                "bearish": bearish,
                "neutral": neutral,
                "reasoning": decision_reasoning,
            }
        )

        agreed: list[str] = []
        disagreed: list[str] = []
        for entry in analyst_entries:
            desc = f"{entry['role']}\uff08{entry['signal_cn']}\uff0c{entry['confidence_text']}\uff09"
            if entry["signal"] == majority:
                agreed.append(desc)
            else:
                disagreed.append(desc)

        conflict_records.append(
            {
                "ticker": ticker,
                "stock_name": stock_name,
                "majority_signal": majority_cn,
                "agreed": agreed,
                "disagreed": disagreed,
            }
        )

        lines.extend(
            [
                _h2(f"标的：{stock_name}（{ticker}）"),
                "",
            ]
        )

        lines.append(_h2("七位大师详细观点"))
        if analyst_entries:
            for idx, entry in enumerate(analyst_entries, start=1):
                lines.extend(
                    [
                        _h4(f"{idx}. {entry['role']}"),
                        f"\u4fe1\u53f7\uff1a{entry['signal_cn']}",
                        f"\u4fe1\u5fc3\u6307\u6570\uff1a{entry['confidence_text']}",
                        "\u6838\u5fc3\u7ed3\u8bba\uff1a",
                    ]
                )
                _append_quote_block(lines, entry["conclusion"])
                lines.append("\u6838\u5fc3\u4f9d\u636e\uff1a")
                _append_quote_block(lines, entry["basis"])
                lines.append(f"\u98ce\u9669\u63d0\u793a\uff1a{entry['risk'] or '\u65e0'}")
                lines.append(f"\u5931\u6548\u89e6\u53d1\uff1a{entry['trigger'] or '\u65e0'}")
                lines.append(f"\u8bc1\u636eID\uff1a{entry['evidence_text']}")
                lines.append("")
        else:
            lines.extend(["\u8bc1\u636e\u4e0d\u8db3\uff0c\u5df2\u5b89\u5168\u964d\u7ea7\u3002", ""])

        lines.extend(
            [
                _h2("总交易决策"),
                f"\u64cd\u4f5c\u5efa\u8bae\uff1a{_action_cn(action)}",
                f"\u80a1\u6570\uff1a{decision.get('quantity', 0)}",
                f"\u4fe1\u5fc3\u6307\u6570\uff1a{confidence_text}",
                "\u51b3\u7b56\u903b\u8f91\uff1a",
            ]
        )
        _append_quote_block(lines, decision_reasoning)
        lines.append("")

    lines.extend(
        [
            _h2("风格隔离校验结果"),
            f"\u68c0\u7d22\u8c03\u7528\u603b\u6570\uff1a{audit['total_calls']}",
            f"\u8de8\u5927\u5e08\u547d\u4e2d\u4e22\u5f03\u4e8b\u4ef6\uff1a{audit['cross_master_drop_events']}",
            f"\u662f\u5426\u53d1\u751f\u98ce\u683c\u6c61\u67d3\uff1a{'\u5426' if audit['is_clean'] else '\u662f'}",
            "",
        ]
    )
    if not audit["is_clean"]:
        lines.append("\u98ce\u683c\u6c61\u67d3\u544a\u8b66\u660e\u7ec6\uff1a")
        for event in audit["events"]:
            lines.append(
                f"- Agent={event.get('agent')}\uff0cMaster={event.get('master')}\uff0c"
                f"\u4e22\u5f03ID={','.join(event.get('dropped_cross_master_ids') or []) or '\u65e0'}\uff0c"
                f"CallID={event.get('call_id')}"
            )
        lines.append("")

    lines.extend(
        [
            _h2("总分析师综合决策说明"),
            "\u603b\u5206\u6790\u5e08\u4f1a\u5148\u68c0\u7d22 chief_analyst memory\uff08\u7531\u4e03\u4f4d\u5927\u5e08"
            "\u672c\u8f6e\u7ed3\u6784\u5316\u8f93\u51fa\u7ec4\u6210\uff09\uff0c\u518d\u7ed9\u51fa\u7ec4\u5408\u51b3\u7b56\u3002",
            "",
            _h2("决策冲突与处理逻辑"),
        ]
    )
    if conflict_records:
        for item in conflict_records:
            lines.extend(
                [
                    _h2(f"{item['stock_name']}（{item['ticker']}）"),
                    f"\u591a\u6570\u6d3e\u4fe1\u53f7\uff1a{item['majority_signal']}",
                    f"\u4e00\u81f4\u89c2\u70b9\uff1a{'\uff1b'.join(item['agreed']) if item['agreed'] else '\u65e0'}",
                    f"\u5206\u6b67\u89c2\u70b9\uff1a{'\uff1b'.join(item['disagreed']) if item['disagreed'] else '\u65e0'}",
                    "\u5904\u7406\u903b\u8f91\uff1a\u591a\u6570\u6d3e\u51b3\u5b9a\u4e3b\u65b9\u5411\uff0c"
                    "\u5206\u6b67\u89c2\u70b9\u7528\u4e8e\u7ea6\u675f\u4ed3\u4f4d\u548c\u98ce\u9669\u66b4\u9732\u3002",
                    "",
                ]
            )
    else:
        lines.extend(["\u65e0\u51b2\u7a81\u6570\u636e\u3002", ""])

    lines.extend([_h2("投资组合汇总"), ""])
    if portfolio_rows:
        for idx, row in enumerate(portfolio_rows, start=1):
            lines.append(
                f"{idx}. {row['ticker']}\uff0c{row['action']}\uff0c{row['quantity']} \u80a1\uff0c"
                f"\u4fe1\u5fc3 {row['confidence_text']}\uff0c"
                f"\u770b\u591a/\u770b\u7a7a/\u4e2d\u6027={row['bullish']}/{row['bearish']}/{row['neutral']}"
            )
    else:
        lines.append("\u6682\u65e0\u7ec4\u5408\u6c47\u603b\u6570\u636e\u3002")
    lines.append("")

    unique_risks = [x for x in dict.fromkeys(item.strip() for item in risk_items if item and item.strip())]
    unique_triggers = [x for x in dict.fromkeys(item.strip() for item in invalidation_items if item and item.strip())]

    lines.append(_h2("风险清单与失效条件"))
    lines.append("")
    lines.append(_h2("风险清单"))
    if unique_risks:
        for idx, item in enumerate(unique_risks[:12], start=1):
            lines.append(f"{idx}. {item}")
    else:
        lines.append("1. \u6682\u65e0")
    lines.append("")

    lines.append(_h2("失效条件"))
    if unique_triggers:
        for idx, item in enumerate(unique_triggers[:12], start=1):
            lines.append(f"{idx}. {item}")
    else:
        lines.append("1. \u6682\u65e0")
    lines.append("")

    lines.append(_h2("可执行动作清单"))
    if portfolio_rows:
        for idx, row in enumerate(portfolio_rows, start=1):
            reason_short = _normalize_reasoning(row["reasoning"])[:220] if row["reasoning"] else "\u65e0"
            lines.append(
                f"{idx}. {row['ticker']}\uff1a{row['action']} {row['quantity']} \u80a1\u3002"
                f"\u7406\u7531\uff1a{reason_short}"
            )
    else:
        lines.append("\u6682\u65e0\u53ef\u6267\u884c\u52a8\u4f5c\u3002")
    lines.append("")

    return "\n".join(lines).replace("#", "＃")


_DIV_LINE_RE = re.compile(r'^<div style="(?P<style>[^"]*)">(?P<text>.*)</div>$')


def _apply_docx_font(run, size_pt: float, *, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "宋体"
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    except Exception:
        pass


def _write_docx_report(content: str, output_path: Path) -> None:
    doc = Document()
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        match = _DIV_LINE_RE.match(line)
        if match:
            style = match.group("style")
            text = unescape(match.group("text")).replace("#", "")
            paragraph = doc.add_paragraph()
            if "text-align:center" in style and "font-size:26pt" in style:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                _apply_docx_font(run, 26, bold=True)
            elif "font-size:14pt" in style:
                run = paragraph.add_run(text)
                _apply_docx_font(run, 14, bold=True)
            else:
                run = paragraph.add_run(text)
                _apply_docx_font(run, 15, bold=True)
            continue

        if line.startswith(">"):
            quote_text = line[1:].strip().replace("#", "")
            paragraph = doc.add_paragraph(quote_text or "（无）")
            paragraph.paragraph_format.left_indent = Pt(18)
            if paragraph.runs:
                _apply_docx_font(paragraph.runs[0], 12)
            continue

        paragraph = doc.add_paragraph(line.replace("#", ""))
        if paragraph.runs:
            _apply_docx_font(paragraph.runs[0], 12)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def _update_overall_statistics(report_path: Path, result: dict) -> None:
    default_stats = _resolve_runtime_path("MASTER_LIBRARY_ROOT", "AgentLibrary") / "overall_statistics.md"
    stats_path = _resolve_runtime_path("MASTER_STATS_PATH", str(default_stats))
    retrieval_logs = result.get("retrieval_logs", []) or []
    audit = isolation_audit(retrieval_logs)
    lines = [
        "",
        "## Latest Fusion Run",
        f"- report_path: {_safe_display_path(report_path)}",
        f"- run_time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- retrieval_calls: {audit['total_calls']}",
        f"- cross_master_drop_events: {audit['cross_master_drop_events']}",
        f"- style_pollution: {'no' if audit['is_clean'] else 'yes'}",
    ]
    try:
        old = stats_path.read_text(encoding="utf-8") if stats_path.exists() else ""
        stats_path.write_text(old + "\n" + "\n".join(lines) + "\n", encoding="utf-8", newline="\n")
    except Exception:
        pass


def export_consolidated_report_to_desktop(result: dict, tickers: list[str], start_date: str, end_date: str) -> dict[str, Path | None]:
    desktop_dir = _get_desktop_dir()
    outputs_dir = _resolve_runtime_path("OUTPUT_DIR", "./outputs")
    outputs_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    primary_report = outputs_dir / f"{stamp}_\u7efc\u5408\u51b3\u7b56\u62a5\u544a.docx"

    content = _build_consolidated_markdown(
        result=result,
        tickers=tickers,
        start_date=start_date,
        end_date=end_date,
    )
    _write_docx_report(content, primary_report)

    desktop_report: Path | None = None
    if desktop_dir is not None:
        desktop_dir.mkdir(parents=True, exist_ok=True)
        if len(tickers) == 1:
            stock_name = _safe_file_name(_resolve_stock_name(tickers[0]))
            desktop_report = desktop_dir / f"{stock_name}_\u7efc\u5408\u7814\u62a5.docx"
        else:
            desktop_report = desktop_dir / "\u7efc\u5408.docx"
        _write_docx_report(content, desktop_report)

    _update_overall_statistics(primary_report, result)
    return {"primary": primary_report, "desktop": desktop_report}


if __name__ == "__main__":
    inputs = parse_cli_inputs(
        description="\u8fd0\u884c A \u80a1 AI \u6295\u7814\u7cfb\u7edf",
        require_tickers=False,
        default_months_back=None,
        include_graph_flag=True,
        include_reasoning_flag=True,
    )

    tickers = inputs.tickers
    portfolio = {
        "cash": inputs.initial_cash,
        "margin_requirement": inputs.margin_requirement,
        "margin_used": 0.0,
        "positions": {
            ticker: {
                "long": 0,
                "short": 0,
                "long_cost_basis": 0.0,
                "short_cost_basis": 0.0,
                "short_margin_used": 0.0,
            }
            for ticker in tickers
        },
        "realized_gains": {ticker: {"long": 0.0, "short": 0.0} for ticker in tickers},
    }

    result = run_hedge_fund(
        tickers=tickers,
        start_date=inputs.start_date,
        end_date=inputs.end_date,
        portfolio=portfolio,
        show_reasoning=inputs.show_reasoning,
        selected_analysts=inputs.selected_analysts,
        model_name=inputs.model_name,
        model_provider=inputs.model_provider,
    )

    print_trading_output(result)
    report_paths = export_consolidated_report_to_desktop(
        result=result,
        tickers=tickers,
        start_date=inputs.start_date,
        end_date=inputs.end_date,
    )
    print(
        f"\n{Fore.GREEN}\u5df2\u5bfc\u51fa\u7efc\u5408\u62a5\u544a\uff1a"
        f"{_safe_display_path(report_paths['primary'])}{Style.RESET_ALL}"
    )
    if report_paths.get("desktop") is not None:
        desktop_path = report_paths["desktop"]
        print(f"{Fore.GREEN}\u53ef\u9009\u684c\u9762\u526f\u672c\uff1a{_safe_display_path(desktop_path)}{Style.RESET_ALL}")
    else:
        print(f"{Fore.YELLOW}\u672a\u914d\u7f6e DESKTOP_OUTPUT_DIR\uff0c\u5df2\u8df3\u8fc7\u684c\u9762\u526f\u672c\u5bfc\u51fa\u3002{Style.RESET_ALL}")
